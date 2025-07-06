from docx import Document

from .base import DocumentParser
from ..layout_analyser import LayoutAnalyzer
from ..types import BoundingBox, DocumentElement, DocumentStructure, ElementType, ParsedDocument


class PositionalWordParser(DocumentParser):
    """Enhanced Word parser that preserves document structure"""
    
    def can_parse(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def parse(self, file_path: str) -> ParsedDocument:
        elements = []
        metadata = {}
        errors = []
        page_dimensions = {}
        
        try:
            if file_path.lower().endswith('.docx'):
                doc = Document(file_path)
                
                metadata = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                }
                
                y_position = 0
                page_num = 0
                
                # Process paragraphs with structure preservation
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Determine element type based on style
                        element_type = self._classify_paragraph(para)
                        
                        # Estimate position (Word doesn't provide exact coordinates)
                        bbox = BoundingBox(
                            x0=0,
                            y0=y_position,
                            x1=600,  # Estimated page width
                            y1=y_position + 20,  # Estimated line height
                            page=page_num
                        )
                        
                        element = DocumentElement(
                            content=para.text,
                            element_type=element_type,
                            bbox=bbox,
                            style={
                                'style_name': para.style.name,
                                'alignment': str(para.alignment) if para.alignment else 'LEFT'
                            }
                        )
                        elements.append(element)
                        y_position += 25
                
                # Process tables with structure
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    table_content = "\n".join(["\t".join(row) for row in table_data])
                    
                    bbox = BoundingBox(
                        x0=0,
                        y0=y_position,
                        x1=600,
                        y1=y_position + len(table_data) * 20,
                        page=page_num
                    )
                    
                    element = DocumentElement(
                        content=table_content,
                        element_type=ElementType.TABLE,
                        bbox=bbox,
                        metadata={'table_data': table_data, 'table_index': table_idx}
                    )
                    elements.append(element)
                    y_position += len(table_data) * 25
                
                page_dimensions[0] = (600, y_position)
                
        except Exception as e:
            errors.append(f"Word parsing error: {str(e)}")
        
        reading_order = LayoutAnalyzer.determine_reading_order(elements)
        
        structure = DocumentStructure(
            elements=elements,
            reading_order=reading_order,
            page_dimensions=page_dimensions
        )
        
        return ParsedDocument(
            structure=structure,
            metadata=metadata,
            file_path=file_path,
            file_type="word",
            parser_used="PositionalWordParser",
            errors=errors if errors else None
        )
    
    def _classify_paragraph(self, paragraph) -> ElementType:
        """Classify paragraph based on style"""
        style_name = paragraph.style.name.lower()
        
        if 'heading' in style_name or 'title' in style_name:
            return ElementType.HEADING
        elif 'caption' in style_name:
            return ElementType.CAPTION
        elif 'list' in style_name:
            return ElementType.LIST
        else:
            return ElementType.PARAGRAPH